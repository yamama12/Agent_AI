from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date, datetime
import os


def generate_CertificatScolarite_docx(data: dict, output_path: str = None) -> str:

    os.makedirs("files", exist_ok=True)
    
    if output_path:
        dirpath = os.path.dirname(output_path)
        if dirpath:
            os.makedirs(dirpath, exist_ok=True)
        filename = os.path.basename(output_path)
    else:
        os.makedirs("files", exist_ok=True)
        filename = f"certificat_scolarite_{data['Matricule']}.docx"
        output_path = f"files/{filename}"
    
    doc = Document()
    
    PRIMARY_COLOR = RGBColor(26, 54, 93)       # #1a365d
    SECONDARY_COLOR = RGBColor(36, 40, 47)     # #24282f
    ACCENT_COLOR = RGBColor(212, 175, 55)      # #d4af37
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    

    def format_date(date_value):
        """Convertit une date en chaîne de caractères"""
        if isinstance(date_value, (date, datetime)):
            return date_value.strftime('%d/%m/%Y')
        elif isinstance(date_value, str):
            try:
                for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y'):
                    try:
                        dt = datetime.strptime(date_value, fmt)
                        return dt.strftime('%d/%m/%Y')
                    except:
                        continue
            except:
                pass
            return date_value
        else:
            return str(date_value) if date_value else 'Non renseigné'
    
    def safe_get(data_dict, key, default='Non renseigné'):
        """Récupère une valeur avec une valeur par défaut"""
        value = data_dict.get(key)
        if value is None:
            return default
        return value

    # EN-TÊTE PRINCIPAL
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = False
    
    header_table.columns[0].width = Cm(3.81)    # Colonne Logo
    header_table.columns[1].width = Cm(12.16)   # Colonne Titre
    header_table.columns[2].width = Cm(3.8)     # Colonne Texte arabe
    
    logo_cell = header_table.cell(0, 0)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    logo_path = "assets/logo_ise.png"
    if os.path.exists(logo_path):
        try:
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(logo_path, width=Cm(5))
        except:
            pass
    
    title_cell = header_table.cell(0, 1)
    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title_run = title_paragraph.add_run("INTERNATIONAL SCHOOL OF ELITE\n")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR
    
    # Sous-titre
    subtitle_run = title_paragraph.add_run("Collège & Lycée ISE")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = SECONDARY_COLOR
    
    # TEXTE ARABE ---
    arabic_cell = header_table.cell(0, 2)
    arabic_paragraph = arabic_cell.paragraphs[0]
    arabic_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    arabic_text = "الجهوية التونسية\nوزارة التربية\nالمندوبية الجهوية للتربية بنابل\nالمدرسة الدولية للنخبة"
    arabic_run = arabic_paragraph.add_run(arabic_text)
    arabic_run.font.size = Pt(8)
    arabic_run.font.bold = True
    arabic_run.font.color.rgb = PRIMARY_COLOR
    
    # SÉPARATION VISUELLE
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    line_paragraph = doc.add_paragraph()
    line_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_paragraph.paragraph_format.space_before = Pt(0)
    line_paragraph.paragraph_format.space_after = Pt(0)
    
    line_run = line_paragraph.add_run("─" * 60)
    line_run.font.color.rgb = ACCENT_COLOR
    line_run.font.size = Pt(10)
    
    # TITRE PRINCIPAL DU CERTIFICAT
    cert_title = doc.add_paragraph()
    cert_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cert_title.paragraph_format.space_before = Pt(6)
    cert_title.paragraph_format.space_after = Pt(6)
    
    cert_run = cert_title.add_run("CERTIFICAT DE SCOLARITÉ")
    cert_run.font.size = Pt(16)
    cert_run.font.bold = True
    cert_run.font.color.rgb = PRIMARY_COLOR
    
    # Année scolaire
    year_paragraph = doc.add_paragraph()
    year_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year_paragraph.paragraph_format.space_after = Pt(16)
    
    year_run = year_paragraph.add_run(f"Année Scolaire {safe_get(data, 'AnneeScolaire', '2024-2025')}")
    year_run.font.size = Pt(10)
    year_run.font.italic = True
    year_run.font.color.rgb = SECONDARY_COLOR
    
    # SECTION 1 : INFORMATIONS DE L'ÉLÈVE
    student_title = doc.add_paragraph()
    student_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    student_title.paragraph_format.space_before = Pt(8)
    student_title.paragraph_format.space_after = Pt(8)
    
    student_title_run = student_title.add_run("Informations de l'élève")
    student_title_run.font.size = Pt(14)
    student_title_run.font.bold = True
    student_title_run.font.underline = True
    student_title_run.font.color.rgb = PRIMARY_COLOR
    
    # Contenu des informations
    student_info = doc.add_paragraph()
    student_info.paragraph_format.space_after = Pt(12)
    student_info.paragraph_format.line_spacing = 1.0
    
    # Liste des informations
    student_info.add_run(f"Nom et Prénom: {safe_get(data, 'NomFr', '')} {safe_get(data, 'PrenomFr', '')}\n").font.size = Pt(12)
    student_info.add_run(f"Date de naissance: {format_date(safe_get(data, 'DateNaissance'))}\n").font.size = Pt(12)
    student_info.add_run(f"Lieu de naissance: {safe_get(data, 'LieuNaissance')}\n").font.size = Pt(12)
    student_info.add_run(f"Nationalité: {safe_get(data, 'Nationalite', 'Tunisienne')}\n").font.size = Pt(12)
    student_info.add_run(f"Téléphone: {safe_get(data, 'Tel1')}\n").font.size = Pt(12)
    student_info.add_run(f"Adresse: {safe_get(data, 'AdresseFr')}\n").font.size = Pt(12)
    
    # SECTION 2 : INFORMATIONS SCOLAIRES
    school_info_title = doc.add_paragraph()
    school_info_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    school_info_title.paragraph_format.space_before = Pt(8)
    school_info_title.paragraph_format.space_after = Pt(8)
    
    school_title_run = school_info_title.add_run("Informations scolaires")
    school_title_run.font.size = Pt(14)
    school_title_run.font.bold = True
    school_title_run.font.underline = True
    school_title_run.font.color.rgb = PRIMARY_COLOR
    
    # Contenu des informations scolaires
    school_info = doc.add_paragraph()
    school_info.paragraph_format.space_after = Pt(16)
    
    school_info.add_run(f"Classe: {safe_get(data, 'Classe')}\n").font.size = Pt(12)
    school_info.add_run(f"Année scolaire: {safe_get(data, 'AnneeScolaire', '2024-2025')}").font.size = Pt(12)
    
    # SECTION 3 : CERTIFICATION
    certification = doc.add_paragraph()
    certification.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    certification.paragraph_format.space_before = Pt(12)
    certification.paragraph_format.space_after = Pt(12)
    
    cert_text = (
        "Je soussignée, Mme Balkis Zrelli, Directrice du Collège et Lycée ISE, "
        "certifie que l'élève mentionné ci-dessus est inscrit(e) et suit une scolarité "
        "régulière dans notre établissement pour l'année scolaire indiquée. "
        "Ce certificat est délivré à la demande des parents/tuteurs."
    )
    certification.add_run(cert_text).font.size = Pt(10)
    
    # SECTION 4 : SIGNATURE
    # Date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.paragraph_format.space_after = Pt(20)
    
    date_run = date_paragraph.add_run(f"Fait à Nabeul, le {date.today().strftime('%d/%m/%Y')}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = SECONDARY_COLOR
    
    # Zone de signature
    signature_box = doc.add_paragraph()
    signature_box.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Ligne de signature
    signature_box.add_run("_________________________\n").font.size = Pt(9)
    
    # Texte signature
    sig_run = signature_box.add_run("Signature et Cachet\n")
    sig_run.font.size = Pt(9)
    sig_run.font.italic = True
    sig_run.font.color.rgb = PRIMARY_COLOR
    
    # PIED DE PAGE
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_text = (
        "INTERNATIONAL SCHOOL OF ELITE - Avenue Mohamed V, Nabeul 8000, Tunisia • "
        "Tél: (+216) 99 555 222 • Email: contact@ise-college-lycee.com"
    )
    
    footer_run = footer_paragraph.add_run(footer_text)
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = PRIMARY_COLOR
    
    # SAUVEGARDE DU DOCUMENT
    doc.save(output_path)
    return filename

